"""CV text extraction for PDF, DOCX and plain-text files."""

from __future__ import annotations

import io
from dataclasses import dataclass

from docx import Document
from pdfminer.high_level import extract_text as pdf_extract_text


class UnsupportedFileError(ValueError):
    """Raised when an uploaded file type cannot be parsed."""


class EmptyDocumentError(ValueError):
    """Raised when a file parses to little or no usable text."""


@dataclass(frozen=True)
class ParsedCV:
    text: str
    filename: str
    file_type: str
    word_count: int


_MIN_WORDS = 30


def _extract_pdf(data: bytes) -> str:
    return pdf_extract_text(io.BytesIO(data)) or ""


def _extract_docx(data: bytes) -> str:
    document = Document(io.BytesIO(data))
    parts: list[str] = [p.text for p in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            parts.extend(cell.text for cell in row.cells)
    return "\n".join(parts)


def _extract_txt(data: bytes) -> str:
    for encoding in ("utf-8", "utf-16", "latin-1"):
        try:
            return data.decode(encoding)
        except UnicodeDecodeError:
            continue
    return data.decode("utf-8", errors="ignore")


def _detect_type(filename: str) -> str:
    lower = filename.lower()
    if lower.endswith(".pdf"):
        return "pdf"
    if lower.endswith(".docx"):
        return "docx"
    if lower.endswith((".txt", ".text", ".md")):
        return "txt"
    raise UnsupportedFileError(
        "Unsupported file type. Please upload a PDF, DOCX or TXT file."
    )


def parse_cv(data: bytes, filename: str) -> ParsedCV:
    """Extract plain text from an uploaded CV.

    Raises ``UnsupportedFileError`` for unknown extensions and
    ``EmptyDocumentError`` when the document yields too little text (often a
    sign of a scanned/image-only PDF that an ATS could not read either).
    """
    file_type = _detect_type(filename)
    if file_type == "pdf":
        text = _extract_pdf(data)
    elif file_type == "docx":
        text = _extract_docx(data)
    else:
        text = _extract_txt(data)

    text = text.replace("\r\n", "\n").replace("\r", "\n").strip()
    word_count = len(text.split())
    if word_count < _MIN_WORDS:
        raise EmptyDocumentError(
            "We couldn't extract readable text from this file. If it's a scanned "
            "or image-based PDF, an ATS won't be able to read it either — export "
            "a text-based PDF or upload a DOCX/TXT version."
        )
    return ParsedCV(
        text=text,
        filename=filename,
        file_type=file_type,
        word_count=word_count,
    )
